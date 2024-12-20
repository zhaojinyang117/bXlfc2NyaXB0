from docx import Document
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import docx2pdf


def data_report():
    # 输入需要修改的内容
    project_address = input("请输入项目地址：")
    contact_person = input("请输入联系人：")
    sampling_date = input("请输入采样日期（格式：MM-DD）：")
    sampling_temperature = input("请输入现场采样温度（℃）：")
    sampling_humidity = input("请输入现场采样湿度（%RH）：")

    # 从采样日期中提取月和日
    sampling_date_obj = datetime.datetime.strptime(sampling_date, "%m-%d")
    month = sampling_date_obj.month
    day = sampling_date_obj.day

    # 判断检测类型
    detection_type = int(input("请确认检测类型：\n->初检：1\n->复检：2\n=>"))
    if detection_type == 1:
        dt = "初检"
    elif detection_type == 2:
        dt = "复检"
    else:
        print("请重新输入")
        detection_type = int(input("请确认检测类型：\n->初检：1\n->复检：2\n->"))
        if detection_type == 1:
            dt = "初检"
        elif detection_type == 2:
            dt = "复检"

    # 加载docx文件
    doc = Document("./resource/模板.docx")  # 假设模板文件名为“模板.docx”

    # 修改第一页的内容
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "<地址>" in run.text:
                run.clear()  # 清空run内容
                run.add_text(project_address)
            if "<联系人>" in run.text:
                run.clear()
                run.add_text(contact_person)
        if "<月>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<月>", str(month))
        if "<日>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<日>", str(day))

    # 修改第三页委托概况的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "<联系人>" in cell.text:
                    cell.text = cell.text.replace("<联系人>", contact_person)
                if "<地址>" in cell.text:
                    cell.text = cell.text.replace("<地址>", project_address)
                if (
                    "Sampling date" in cell.text
                    and "<月>" in cell.text
                    and "<日>" in cell.text
                ):
                    cell.text = f"{sampling_date_obj.year} 年 {month} 月 {day} 日"

    # 修改环境条件的内容
    for paragraph in doc.paragraphs:
        if "<温度>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<温度>", sampling_temperature)
        if "<湿度>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<湿度>", sampling_humidity)

    # 填写第四页检测结果
    for table in doc.tables:
        if table.cell(0, 0).text == "序":
            # 处理列（点位）和对应的值
            for i in range(4):
                # 获取点位输入
                point_input = input(f"点位{i+1}:")
                
                # 处理点位
                cell_point = table.cell(i+3, 1)
                for paragraph in cell_point.paragraphs:
                    run = paragraph.add_run()
                    run.font.size = Pt(10.5)
                    run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.text = point_input

                # 处理值
                cell_value = table.cell(i+3, 2)
                for paragraph in cell_value.paragraphs:
                    run = paragraph.add_run()
                    run.font.size = Pt(10.5)
                    run.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run.text = input(f"值{i+1}:")

                # 如果有输入点位，在第三列填入"≤0.08"
                if point_input.strip():  # 检查点位是否有实际输入（排除空格）
                    cell_limit = table.cell(i+3, 3)
                    for paragraph in cell_limit.paragraphs:
                        run = paragraph.add_run()
                        run.font.size = Pt(10.5)
                        run.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.text = "≤0.08"

    # 保存修改后的文件
    doc.save(
        f"../{project_address}+{dt}报告+{month}.{day}.docx"
    )  # 保存为“修改后的模板.docx”

    # 转成pdf
    docx2pdf.convert(f"../{project_address}+{dt}报告+{month}.{day}.docx")
